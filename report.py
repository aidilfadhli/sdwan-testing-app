"""Generate Berita Acara Uji Fungsi (.docx) dari template dokumen asli (revisi Jul 2026).

Pemetaan tabel pada ba_template.docx:
  tabel 0 : identitas — Type Device, S/N, Lokasi, Tanggal, Petugas (kolom ke-3 diisi nilai)
  tabel 1 : checklist 6 item (kolom Hasil & Keterangan)
  tabel 2-6 : kotak foto (fisik, led, fortios, interface, ping)
  tabel 7 : catatan tambahan
  tabel 8 : tanda tangan — baris 0: Petugas Penguji; baris 1: Mengetahui/Menyetujui (3 nama)
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from checklist import VENDOR_REGISTRY, get_vendor_spec, CHECKLIST_ITEMS, PHOTO_SECTIONS

BASE_DIR = Path(__file__).resolve().parent

GREEN = RGBColor(0x1E, 0x7D, 0x32)
RED = RGBColor(0xC6, 0x28, 0x28)


def _set_cell_text(cell, text, bold=False, color=None, align=None):
    cell.text = ""
    para = cell.paragraphs[0]
    if align is not None:
        para.alignment = align
    run = para.add_run(text or "")
    run.font.size = Pt(10)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def _fill_photos(table, photo_paths):
    cell = table.rows[0].cells[0]
    if not photo_paths:
        _set_cell_text(cell, "( Tidak ada foto )", align=WD_ALIGN_PARAGRAPH.CENTER)
        return
    cell.text = ""
    for i, path in enumerate(photo_paths):
        para = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        try:
            run.add_picture(str(path), width=Inches(4.8))
        except Exception:
            para.add_run(f"( Gagal memuat foto: {Path(path).name} )")


def _replace_para_text(para, text):
    for run in para.runs:
        run.text = ""
    if para.runs:
        para.runs[0].text = text
    else:
        para.add_run(text)


def _fill_signatures(table, petugas, approvers):
    """Mengisi tabel tanda tangan 2x2 grid:
    - Row 0, Cell 0: Petugas Penguji
    - Row 0, Cell 1: Mengetahui / Saksi (DCS)
    - Row 1, Cell 0: Mengetahui / Saksi (DID)
    - Row 1, Cell 1: Mengetahui / Saksi (ECS)
    """
    if not table.rows:
        return

    def fill_cell(cell, name):
        for para in cell.paragraphs:
            if para.text.strip().startswith("(_"):
                _replace_para_text(para, f"(  {name}  )" if name else "(________________________)")
            elif "Nama / NIP" in para.text:
                _replace_para_text(para, "")

    # Row 0: Petugas Penguji & Saksi 1 (DCS)
    cells0 = table.rows[0].cells
    if len(cells0) > 0:
        fill_cell(cells0[0], petugas)
    if len(cells0) > 1:
        fill_cell(cells0[1], approvers[0] if len(approvers) > 0 else "")

    # Row 1: Saksi 2 (DID) & Saksi 3 (ECS)
    if len(table.rows) > 1:
        cells1 = table.rows[1].cells
        if len(cells1) > 0:
            fill_cell(cells1[0], approvers[1] if len(approvers) > 1 else "")
        if len(cells1) > 1:
            fill_cell(cells1[1], approvers[2] if len(approvers) > 2 else "")


def generate_ba(report: dict, photos_by_section: dict, out_path: Path) -> Path:
    """report: dict baris tabel reports; photos_by_section: {section: [Path, ...]}"""
    vendor_id = report.get("vendor") or "fortinet"
    vendor_spec = get_vendor_spec(vendor_id)
    
    template_file = vendor_spec.get("template_file", "ba_template_fortinet.docx")
    template_path = BASE_DIR / template_file
    if not template_path.exists():
        # Fallback if specific file missing
        template_path = BASE_DIR / "ba_template_fortinet.docx"
        if not template_path.exists():
            template_path = BASE_DIR / "ba_template.docx"

    doc = Document(template_path)

    ident = doc.tables[0]
    values = [
        report.get("type_device", ""),
        report.get("serial_number", ""),
        "-", # Nomor Aset / BA
        report.get("lokasi", ""),
        report.get("tanggal", ""),
        report.get("petugas", ""),
    ]
    for row, val in zip(ident.rows, values):
        _set_cell_text(row.cells[2], val, bold=True)

    check = doc.tables[1]
    items = vendor_spec.get("items", CHECKLIST_ITEMS)
    for i, item in enumerate(items, start=1):
        if i >= len(check.rows):
            break
        hasil = report.get(f"hasil{i}", "")
        ket = report.get(f"ket{i}", "")
        color = GREEN if hasil == "OK" else (RED if hasil == "NOT OK" else None)
        _set_cell_text(check.rows[i].cells[3], hasil, bold=True, color=color,
                       align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_text(check.rows[i].cells[4], ket)

    sections = vendor_spec.get("photo_sections", PHOTO_SECTIONS)
    for offset, section in enumerate(sections):
        if (2 + offset) < len(doc.tables) - 2:
            _fill_photos(doc.tables[2 + offset], photos_by_section.get(section["key"], []))

    # Catatan Tambahan (Table 8) & Tanda Tangan (Table 9)
    catatan_table_idx = len(doc.tables) - 2
    sig_table_idx = len(doc.tables) - 1
    
    if catatan_table_idx >= 0:
        _set_cell_text(doc.tables[catatan_table_idx].rows[0].cells[0], report.get("catatan", ""))

    if sig_table_idx >= 0:
        _fill_signatures(
            doc.tables[sig_table_idx],
            report.get("petugas", ""),
            [report.get("saksi", ""), report.get("saksi2", ""), report.get("saksi3", "")],
        )

    doc.save(out_path)
    return out_path

